"""
Servicio de Generación de Documentos
Genera documentos Word a partir de plantillas con datos de pólizas y los convierte a PDF
"""
from pathlib import Path
from datetime import datetime, timedelta
from typing import Dict, Any, Tuple
from docx import Document
import calendar
import platform
import subprocess


class GenerateDocumentService:
    """
    Servicio para generar documentos Word de pólizas
    """
    
    def __init__(self):
        """Inicializa el servicio con las rutas de plantillas"""
        self.base_path = Path(__file__).parent.parent.parent
        self.plantilla_path = self.base_path / "assets" / "condicionado_particular" / "rumbo_plantilla_sme.docx"
        self.output_path = self.base_path / "db" / "documentos"
        
        # Crear directorio de salida si no existe
        self.output_path.mkdir(parents=True, exist_ok=True)
    
    def obtener_ultimo_dia_mes(self, fecha: datetime) -> datetime:
        """
        Obtiene el último día del mes de una fecha dada
        
        Args:
            fecha: Fecha de referencia
            
        Returns:
            datetime: Fecha del último día del mes
        """
        ultimo_dia = calendar.monthrange(fecha.year, fecha.month)[1]
        return fecha.replace(day=ultimo_dia)
    
    def formatear_fecha(self, fecha: datetime, formato: str = "%d/%m/%Y") -> str:
        """
        Formatea una fecha según el formato especificado
        
        Args:
            fecha: Fecha a formatear
            formato: Formato de salida
            
        Returns:
            str: Fecha formateada
        """
        return fecha.strftime(formato)
    
    def formatear_fecha_hora(self, fecha: datetime, formato: str = "%d/%m/%Y %H:%M:%S") -> str:
        """
        Formatea una fecha con hora según el formato especificado
        
        Args:
            fecha: Fecha a formatear
            formato: Formato de salida
            
        Returns:
            str: Fecha y hora formateada
        """
        return fecha.strftime(formato)
    
    def preparar_datos_documento(self, datos_poliza: Dict[str, Any]) -> Dict[str, str]:
        """
        Prepara los datos de la póliza para el documento Word
        Mapea los campos del JSON a los marcadores del documento
        Los marcadores usan formato «marcador» (comillas angulares francesas)
        
        Args:
            datos_poliza: Datos completos de la póliza desde el JSON
            
        Returns:
            Dict: Diccionario con los marcadores y sus valores
        """
        # Parsear fecha de emisión
        fecha_emision = datetime.fromisoformat(datos_poliza["fecha_emision"])
        
        # Generar número de póliza en formato RumbIA###
        id_poliza = datos_poliza["id_poliza"]
        numero_poliza = f"RumbIA{id_poliza:03d}"
        
        # Fecha y hora de inicio de vigencia (fecha emisión a las 00:00:00)
        fecha_vigencia_inicio = fecha_emision.replace(hour=0, minute=0, second=0, microsecond=0)
        
        # Fecha y hora de fin de vigencia (último día del mes a las 23:59:59)
        # Asumiendo vigencia por defecto (puede ajustarse según necesidad)
        fecha_vigencia_fin = self.obtener_ultimo_dia_mes(fecha_emision)
        fecha_vigencia_fin = fecha_vigencia_fin.replace(hour=23, minute=59, second=59, microsecond=0)
        
        # Formatear género
        genero_completo = "Masculino" if datos_poliza["cliente"]["genero"] == "M" else "Femenino"
        
        # Calcular primas y tasas adicionales
        prima_anual = datos_poliza["cotizacion"]["prima_anual"]
        prima_comercial_frecuencia = prima_anual / 10
        # Los seguros están exentos de IGV, por lo tanto primaComercialConIGV = primaComercialXFrecuenciaPago
        prima_comercial_con_igv = prima_comercial_frecuencia
        tasa_implicita = datos_poliza["cotizacion"]["tasa_implicita"]
        
        # Obtener mes en español
        meses_espanol = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        mes_emision = meses_espanol[fecha_emision.month - 1]
        
        # Parsear tabla de devolución (viene como string JSON)
        import json as json_module
        tabla_devolucion_str = datos_poliza["cotizacion"]["tabla_devolucion"]
        tabla_devolucion = json_module.loads(tabla_devolucion_str) if isinstance(tabla_devolucion_str, str) else tabla_devolucion_str
        
        # Números en letras para plazos (hasta 100)
        numeros_letras = {
            1: "uno", 2: "dos", 3: "tres", 4: "cuatro", 5: "cinco",
            6: "seis", 7: "siete", 8: "ocho", 9: "nueve", 10: "diez",
            11: "once", 12: "doce", 13: "trece", 14: "catorce", 15: "quince",
            16: "dieciséis", 17: "diecisiete", 18: "dieciocho", 19: "diecinueve", 20: "veinte",
            30: "treinta", 40: "cuarenta", 50: "cincuenta", 60: "sesenta"
        }
        
        # Calcular plazo de vigencia (número de años de la tabla de devolución)
        plazo_vigencia_anos = len(tabla_devolucion)
        plazo_vigencia_letras = numeros_letras.get(plazo_vigencia_anos, str(plazo_vigencia_anos))
        
        # Mapeo de campos según la especificación
        # IMPORTANTE: Los marcadores en la plantilla usan «marcador» (no {marcador})
        marcadores = {
            # Número de póliza
            "numeroPoliza": numero_poliza,
            
            # Datos del cliente
            "clienteNumeroDocumento": datos_poliza["cliente"]["dni"],
            "clienteNombre": datos_poliza["cliente"]["nombre"],
            "clienteNombreUpper": datos_poliza["cliente"]["nombre"].upper(),
            "clienteFechaNacimiento": datos_poliza["cliente"]["fechaNacimiento"],
            "clienteGenero": genero_completo,
            "clienteTelefono": datos_poliza["cliente"]["telefono"],
            "clienteEmail": datos_poliza["cliente"]["correo"],
            
            # Datos de la cotización
            "clienteEdadActuarial": str(datos_poliza["cotizacion"]["parametros"]["edad_actuarial"]),
            
            # Período de pago de primas (por defecto mensual, puede ajustarse)
            "periodoPagoPrimas": "Mensual",
            
            # Fechas de emisión
            "fechaEmisionPoliza": self.formatear_fecha(fecha_emision),
            "fechaHoraEmisionPoliza": self.formatear_fecha_hora(fecha_emision),
            
            # Fechas de vigencia
            "fechaHoraInicioVigencia": self.formatear_fecha_hora(fecha_vigencia_inicio),
            "fechaHoraFinVigencia": self.formatear_fecha_hora(fecha_vigencia_fin),
            
            # Fechas para firma (día, mes, año por separado)
            "diaEmisionPolizaFirma": str(fecha_emision.day),
            "mesEmisionPolizaFirma": mes_emision,
            "anioEmisionPolizaFirma": str(fecha_emision.year),
            
            # Datos adicionales de la póliza
            "sumaAsegurada": f"S/ {datos_poliza['cotizacion']['suma_asegurada']:,.2f}",
            "primaAnual": f"S/ {datos_poliza['cotizacion']['prima_anual']:,.2f}",
            "primaMensual": f"S/ {datos_poliza['cotizacion']['parametros']['prima']:,.2f}",
            "devolucion": f"S/ {datos_poliza['cotizacion']['devolucion']:,.2f}",
            "producto": datos_poliza["cotizacion"]["producto"],
            "tasaImplicita": f"{datos_poliza['cotizacion']['tasa_implicita'] * 100:.4f}%",
            "porcentajeDevolucion": f"{datos_poliza['cotizacion']['porcentaje_devolucion'] * 100:.2f}%",
            
            # Primas comerciales adicionales
            "tasaAnualCobPrincipal": f"{tasa_implicita * 100:.4f}%",
            "primaComercialAnualPrincipal": f"S/ {prima_anual:,.2f}",
            "primaComercialAnualTotal": f"S/ {prima_anual:,.2f}",
            "primaComercialAnual": f"S/ {prima_anual:,.2f}",
            "primaComercialXFrecuenciaPago": f"S/ {prima_comercial_frecuencia:,.2f}",
            "primaComercialConIGV": f"S/ {prima_comercial_con_igv:,.2f}",
            
            # Plazos
            "plazoVigencia": str(plazo_vigencia_anos),
            "plazoDevolucionAnticipada": str(plazo_vigencia_anos),
            "plazoDevolucionAnticipadaLetras": plazo_vigencia_letras
        }
        
        # Agregar tabla de devolución dinámica (hasta 52 años)
        for i, porcentaje in enumerate(tabla_devolucion, start=1):
            anio = i
            marcadores[f"devolucionAnio{anio}"] = str(anio)
            marcadores[f"devolucionPriPje{anio}"] = f"{porcentaje:.2f}%"
        
        # Rellenar años faltantes hasta 52 con valores vacíos o cero
        for anio in range(len(tabla_devolucion) + 1, 53):
            marcadores[f"devolucionAnio{anio}"] = ""
            marcadores[f"devolucionPriPje{anio}"] = ""
        
        return marcadores
    
    def _reemplazar_en_parrafo(self, paragraph, marcadores: Dict[str, str]):
        """
        Reemplaza marcadores en un párrafo de manera robusta
        Maneja el caso donde los marcadores están fragmentados en múltiples runs
        Los marcadores usan formato «marcador» (comillas angulares francesas)
        
        Args:
            paragraph: Párrafo del documento
            marcadores: Diccionario de marcadores y valores
        """
        # Obtener el texto completo del párrafo
        texto_completo = paragraph.text
        texto_original = texto_completo
        
        # Reemplazar todos los marcadores en el texto completo
        # Los marcadores usan formato «marcador» (comillas angulares)
        for marcador, valor in marcadores.items():
            patron = f"«{marcador}»"
            if patron in texto_completo:
                texto_completo = texto_completo.replace(patron, str(valor))
        
        # Si hubo cambios, actualizar el párrafo
        if texto_completo != texto_original:
            # Limpiar runs existentes y crear uno nuevo con el texto modificado
            # Preservar el formato del primer run
            if paragraph.runs:
                primer_run = paragraph.runs[0]
                # Guardar el formato
                bold = primer_run.bold
                italic = primer_run.italic
                underline = primer_run.underline
                font_name = primer_run.font.name
                font_size = primer_run.font.size
                
                # Limpiar todos los runs
                for run in paragraph.runs:
                    run.text = ""
                
                # Establecer el nuevo texto en el primer run
                primer_run.text = texto_completo
                
                # Restaurar formato
                if bold is not None:
                    primer_run.bold = bold
                if italic is not None:
                    primer_run.italic = italic
                if underline is not None:
                    primer_run.underline = underline
                if font_name:
                    primer_run.font.name = font_name
                if font_size:
                    primer_run.font.size = font_size
            else:
                # Si no hay runs, agregar uno nuevo
                paragraph.add_run(texto_completo)
    
    def reemplazar_marcadores_en_documento(self, doc: Document, marcadores: Dict[str, str]) -> Document:
        """
        Reemplaza los marcadores en el documento Word de manera robusta
        
        Args:
            doc: Documento Word cargado
            marcadores: Diccionario de marcadores y valores
            
        Returns:
            Document: Documento modificado
        """
        # Reemplazar en párrafos principales
        for paragraph in doc.paragraphs:
            self._reemplazar_en_parrafo(paragraph, marcadores)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._reemplazar_en_parrafo(paragraph, marcadores)
        
        # Reemplazar en encabezados y pies de página
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                self._reemplazar_en_parrafo(paragraph, marcadores)
            
            # Footer
            for paragraph in section.footer.paragraphs:
                self._reemplazar_en_parrafo(paragraph, marcadores)
            
            # Tablas en header
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._reemplazar_en_parrafo(paragraph, marcadores)
            
            # Tablas en footer
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._reemplazar_en_parrafo(paragraph, marcadores)
        
        return doc
    
    def convertir_a_pdf(self, ruta_docx: str) -> str:
        """
        Convierte un archivo Word a PDF
        
        Args:
            ruta_docx: Ruta del archivo Word
            
        Returns:
            str: Ruta del archivo PDF generado
        """
        from docx2pdf import convert
        
        # Generar ruta del PDF (mismo nombre, extensión .pdf)
        ruta_pdf = str(Path(ruta_docx).with_suffix('.pdf'))
        
        try:
            # Convertir a PDF
            convert(ruta_docx, ruta_pdf)
            return ruta_pdf
        except Exception as e:
            print(f"⚠️ Error al convertir a PDF: {e}")
            # Si falla la conversión, intentar método alternativo
            return self._convertir_pdf_alternativo(ruta_docx)
    
    def _convertir_pdf_alternativo(self, ruta_docx: str) -> str:
        """
        Método alternativo para convertir Word a PDF usando LibreOffice
        (útil si docx2pdf falla o no está disponible MS Word)
        
        Args:
            ruta_docx: Ruta del archivo Word
            
        Returns:
            str: Ruta del archivo PDF o None si falla
        """
        ruta_pdf = str(Path(ruta_docx).with_suffix('.pdf'))
        
        try:
            # Intentar con LibreOffice (si está instalado)
            if platform.system() == "Windows":
                # Rutas comunes de LibreOffice en Windows
                libreoffice_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                ]
                
                for path in libreoffice_paths:
                    if Path(path).exists():
                        subprocess.run([
                            path,
                            "--headless",
                            "--convert-to", "pdf",
                            "--outdir", str(Path(ruta_docx).parent),
                            ruta_docx
                        ], check=True, capture_output=True)
                        return ruta_pdf
            
            print("⚠️ No se pudo convertir a PDF con LibreOffice")
            return None
        except Exception as e:
            print(f"⚠️ Error en conversión alternativa: {e}")
            return None
    
    def generar_documento(self, datos_poliza: Dict[str, Any], generar_pdf: bool = True, solo_pdf: bool = True) -> Tuple[str, str]:
        """
        Genera el documento Word de la póliza y opcionalmente lo convierte a PDF
        
        Args:
            datos_poliza: Datos completos de la póliza
            generar_pdf: Si se debe generar también el PDF (default: True)
            solo_pdf: Si es True, elimina el Word después de generar el PDF (default: True)
            
        Returns:
            Tuple[str, str]: (ruta_docx, ruta_pdf) - ruta_pdf puede ser None si falla
        """
        # Preparar los datos
        marcadores = self.preparar_datos_documento(datos_poliza)
        
        # Cargar la plantilla
        doc = Document(self.plantilla_path)
        
        # Reemplazar marcadores
        doc = self.reemplazar_marcadores_en_documento(doc, marcadores)
        
        # Generar nombre del archivo de salida
        id_poliza = datos_poliza["id_poliza"]
        nombre_archivo_docx = f"RumbIA{id_poliza:03d}_Condicionado_Particular.docx"
        ruta_docx = self.output_path / nombre_archivo_docx
        
        # Guardar el documento Word
        doc.save(ruta_docx)
        
        ruta_pdf = None
        # Convertir a PDF si está habilitado
        if generar_pdf:
            try:
                ruta_pdf = self.convertir_a_pdf(str(ruta_docx))
                if ruta_pdf:
                    print(f"✅ PDF generado: {ruta_pdf}")
                    
                    # Si solo queremos el PDF, eliminamos el Word
                    if solo_pdf and ruta_pdf:
                        try:
                            import os
                            os.remove(ruta_docx)
                            print(f"🗑️  Archivo Word eliminado: {ruta_docx}")
                        except Exception as e:
                            print(f"⚠️ No se pudo eliminar el Word: {e}")
            except Exception as e:
                print(f"⚠️ No se pudo generar el PDF: {e}")
        
        return str(ruta_docx), ruta_pdf

