"""
Script de prueba para generación de documentos
"""
import json
from pathlib import Path
from app.services import GenerateDocumentService
from docx import Document

def listar_marcadores_en_plantilla():
    """Lista todos los marcadores encontrados en la plantilla"""
    plantilla_path = Path("assets/condicionado_particular/rumbo_plantilla_sme.docx")
    
    if not plantilla_path.exists():
        print(f"❌ Plantilla no encontrada en: {plantilla_path}")
        return
    
    print("🔍 Analizando plantilla...")
    print(f"📄 Ruta: {plantilla_path}\n")
    
    doc = Document(plantilla_path)
    marcadores_encontrados = set()
    
    # Buscar en párrafos
    for i, paragraph in enumerate(doc.paragraphs):
        texto = paragraph.text
        if '«' in texto and '»' in texto:
            print(f"Párrafo {i}: {texto[:100]}...")
            # Extraer marcadores con comillas angulares «marcador»
            import re
            marcadores = re.findall(r'«(\w+)»', texto)
            marcadores_encontrados.update(marcadores)
    
    # Buscar en tablas
    for i_table, table in enumerate(doc.tables):
        for i_row, row in enumerate(table.rows):
            for i_cell, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    texto = paragraph.text
                    if '«' in texto and '»' in texto:
                        print(f"Tabla {i_table}, Fila {i_row}, Celda {i_cell}: {texto[:100]}...")
                        import re
                        marcadores = re.findall(r'«(\w+)»', texto)
                        marcadores_encontrados.update(marcadores)
    
    print("\n📋 Marcadores encontrados en la plantilla:")
    for marcador in sorted(marcadores_encontrados):
        print(f"  • «{marcador}»")
    
    print(f"\n✅ Total: {len(marcadores_encontrados)} marcadores únicos")
    
    return marcadores_encontrados


def probar_generacion_documento():
    """Prueba la generación de un documento con datos de ejemplo"""
    # Cargar póliza de ejemplo
    poliza_path = Path("db/RumbIA001.json")
    
    if not poliza_path.exists():
        print("❌ No se encontró póliza de ejemplo en db/RumbIA001.json")
        return
    
    print("\n📄 Cargando datos de póliza...")
    with open(poliza_path, 'r', encoding='utf-8') as f:
        datos_poliza = json.load(f)
    
    print("✅ Datos cargados")
    print(f"   Cliente: {datos_poliza['cliente']['nombre']}")
    print(f"   ID Póliza: {datos_poliza['id_poliza']}")
    
    # Generar documento
    print("\n🔧 Generando documento...")
    servicio = GenerateDocumentService()
    
    # Preparar marcadores
    marcadores = servicio.preparar_datos_documento(datos_poliza)
    
    print("\n📊 Marcadores preparados para reemplazo:")
    for key, value in list(marcadores.items())[:10]:  # Mostrar primeros 10
        print(f"   «{key}» → {value}")
    print(f"   ... ({len(marcadores)} marcadores en total)")
    
    try:
        ruta_word, ruta_pdf = servicio.generar_documento(datos_poliza, generar_pdf=False)
        print(f"\n✅ Documento generado exitosamente:")
        print(f"   Word: {ruta_word}")
        
        # Verificar contenido del documento generado
        doc_generado = Document(ruta_word)
        print(f"\n🔍 Verificando documento generado...")
        
        tiene_marcadores = False
        for paragraph in doc_generado.paragraphs:
            if '«' in paragraph.text and '»' in paragraph.text:
                tiene_marcadores = True
                print(f"   ⚠️  Marcador sin reemplazar encontrado: {paragraph.text[:100]}...")
        
        if not tiene_marcadores:
            print("   ✅ No se encontraron marcadores sin reemplazar")
        else:
            print("   ❌ Aún hay marcadores sin reemplazar en el documento")
        
    except Exception as e:
        print(f"\n❌ Error al generar documento: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    print("=" * 60)
    print("🧪 TEST DE GENERACIÓN DE DOCUMENTOS")
    print("=" * 60)
    
    # 1. Listar marcadores en plantilla
    print("\n1️⃣  PASO 1: Analizar plantilla")
    print("-" * 60)
    marcadores_plantilla = listar_marcadores_en_plantilla()
    
    # 2. Probar generación
    print("\n2️⃣  PASO 2: Probar generación de documento")
    print("-" * 60)
    probar_generacion_documento()
    
    print("\n" + "=" * 60)
    print("✅ Test completado")
    print("=" * 60)

